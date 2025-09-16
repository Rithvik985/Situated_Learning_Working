"""
Document Service for generating Word documents from assignments
"""

import io
import uuid
from datetime import datetime
from typing import List, Dict, Optional
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from database.models import GeneratedAssignment, AssignmentRubric
import logging

logger = logging.getLogger(__name__)

class DocumentService:
    def __init__(self):
        pass

    def generate_assignment_document(
        self, 
        assignments: List[GeneratedAssignment], 
        rubric: Optional[AssignmentRubric] = None,
        assignment_name: str = "Generated Assignment"
    ) -> io.BytesIO:
        """Generate a Word document containing selected assignments and optional rubric"""
        
        # Create new document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add custom styles
        self._add_custom_styles(doc)
        
        # Add header
        self._add_document_header(doc, assignment_name, assignments[0] if assignments else None)
        
        # Add assignments
        for i, assignment in enumerate(assignments, 1):
            self._add_assignment_section(doc, assignment, i, len(assignments) > 1)
            
            # Add page break between assignments (except for the last one)
            if i < len(assignments):
                doc.add_page_break()
        
        # Add rubric if provided
        if rubric:
            doc.add_page_break()
            self._add_rubric_section(doc, rubric)
        
        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer

    def _add_custom_styles(self, doc: Document):
        """Add custom styles to the document"""
        styles = doc.styles
        
        # Title style
        if 'Assignment Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Assignment Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(18)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Subtitle style
        if 'Assignment Subtitle' not in [s.name for s in styles]:
            subtitle_style = styles.add_style('Assignment Subtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_font = subtitle_style.font
            subtitle_font.name = 'Calibri'
            subtitle_font.size = Pt(14)
            subtitle_font.bold = True
            subtitle_style.paragraph_format.space_after = Pt(8)
        
        # Header style
        if 'Section Header' not in [s.name for s in styles]:
            header_style = styles.add_style('Section Header', WD_STYLE_TYPE.PARAGRAPH)
            header_font = header_style.font
            header_font.name = 'Calibri'
            header_font.size = Pt(12)
            header_font.bold = True
            header_style.paragraph_format.space_before = Pt(12)
            header_style.paragraph_format.space_after = Pt(6)
        
        # Body text style
        if 'Assignment Body' not in [s.name for s in styles]:
            body_style = styles.add_style('Assignment Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)
            body_style.paragraph_format.line_spacing = 1.15
            body_style.paragraph_format.space_after = Pt(6)

    def _add_document_header(self, doc: Document, assignment_name: str, first_assignment: Optional[GeneratedAssignment]):
        """Add document header with assignment information"""
        
        # Main title
        title = doc.add_paragraph(assignment_name, style='Assignment Title')
        
        # Course information
        if first_assignment:
            course_info = doc.add_paragraph()
            course_info.add_run(f"Course: {first_assignment.course_name}").bold = True
            if first_assignment.course_id:
                course_info.add_run(f" ({first_assignment.course_id})")
            
            # Topics and domains
            if first_assignment.topics:
                topics_para = doc.add_paragraph()
                topics_para.add_run("Topics: ").bold = True
                topics_para.add_run(", ".join(first_assignment.topics))
            
            if first_assignment.domains:
                domains_para = doc.add_paragraph()
                domains_para.add_run("Industry Domains: ").bold = True
                domains_para.add_run(", ".join(first_assignment.domains))
        
        # Add separator line
        doc.add_paragraph("_" * 80)

    def _add_assignment_section(self, doc: Document, assignment: GeneratedAssignment, index: int, multiple_assignments: bool):
        """Add an individual assignment section to the document"""
        
        # Assignment title
        if multiple_assignments:
            title_text = f"Assignment {index}: {assignment.title}"
        else:
            title_text = assignment.title
        
        title_para = doc.add_paragraph(title_text, style='Assignment Subtitle')
        
        # Assignment metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run("Difficulty Level: ").bold = True
        meta_para.add_run(assignment.difficulty_level or "Not specified")
        
        if assignment.tags:
            meta_para.add_run(" | Tags: ").bold = True
            meta_para.add_run(", ".join(assignment.tags))
        
        # Assignment description/content
        content_para = doc.add_paragraph(style='Assignment Body')
        
        # Parse and format the assignment description
        lines = assignment.description.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                content_para = doc.add_paragraph(style='Assignment Body')
                continue
            
            # Check if line is a header (starts with #, numbers, or keywords)
            if (line.startswith(('#', '**')) or 
                line.startswith(('Task:', 'Tasks:', 'Assignment:', 'Context:', 'Scenario:')) or
                line.match(r'^\d+\.') if hasattr(line, 'match') else line.startswith(tuple('123456789'))):
                
                # Add as header or bold text
                if content_para.text.strip():  # If current paragraph has content, start new one
                    content_para = doc.add_paragraph(style='Assignment Body')
                
                # Clean up formatting markers
                clean_line = line.replace('**', '').replace('#', '').strip()
                content_para.add_run(clean_line).bold = True
                content_para = doc.add_paragraph(style='Assignment Body')
            else:
                # Add as regular content
                if content_para.text:
                    content_para.add_run(' ')
                content_para.add_run(line)

    def _add_rubric_section(self, doc: Document, rubric: AssignmentRubric):
        """Add rubric section to the document"""
        
        # Rubric title
        rubric_title = doc.add_paragraph("Evaluation Rubric", style='Assignment Title')
        
        # Rubric metadata
        if rubric.rubric_name:
            doc.add_paragraph(rubric.rubric_name, style='Assignment Subtitle')
        
        if rubric.doc_type:
            meta_para = doc.add_paragraph()
            meta_para.add_run("Document Type: ").bold = True
            meta_para.add_run(rubric.doc_type)
        
        # Rubric content
        if isinstance(rubric.criteria, dict) and 'rubrics' in rubric.criteria:
            for category_data in rubric.criteria['rubrics']:
                # Category header
                category_para = doc.add_paragraph(style='Section Header')
                category_para.add_run(category_data.get('category', 'Category')).bold = True
                
                # Category questions
                questions = category_data.get('questions', [])
                for i, question in enumerate(questions, 1):
                    question_para = doc.add_paragraph(style='Assignment Body')
                    question_para.add_run(f"{i}. ").bold = True
                    question_para.add_run(question)
                
                # Add spacing between categories
                doc.add_paragraph()

    def create_download_response(self, doc_buffer: io.BytesIO, filename: str) -> StreamingResponse:
        """Create a StreamingResponse for document download"""
        
        def iter_file():
            yield doc_buffer.getvalue()
        
        # Ensure filename has .docx extension
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        return StreamingResponse(
            iter_file(),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
        )

    def generate_single_assignment_download(
        self, 
        db: Session, 
        assignment_id: str
    ) -> StreamingResponse:
        """Generate download for a single assignment"""
        
        try:
            assignment = db.query(GeneratedAssignment).filter(
                GeneratedAssignment.id == uuid.UUID(assignment_id)
            ).first()
            
            if not assignment:
                raise ValueError("Assignment not found")
            
            # Generate document
            doc_buffer = self.generate_assignment_document(
                assignments=[assignment],
                assignment_name=assignment.title
            )
            
            # Create filename
            safe_title = "".join(c for c in assignment.title if c.isalnum() or c in (' ', '-', '_')).strip()
            filename = f"{safe_title}_Assignment"
            
            return self.create_download_response(doc_buffer, filename)
            
        except Exception as e:
            logger.error(f"Error generating assignment download: {str(e)}")
            raise

    def generate_multiple_assignments_download(
        self, 
        db: Session, 
        assignment_ids: List[str],
        assignment_name: str = "Generated Assignments"
    ) -> StreamingResponse:
        """Generate download for multiple assignments"""
        
        try:
            assignments = db.query(GeneratedAssignment).filter(
                GeneratedAssignment.id.in_([uuid.UUID(aid) for aid in assignment_ids])
            ).all()
            
            if not assignments:
                raise ValueError("No assignments found")
            
            # Get associated rubric if any
            rubric = None
            rubric_query = db.query(AssignmentRubric).filter(
                AssignmentRubric.assignment_ids.contains([assignments[0].id])
            ).first()
            
            if rubric_query:
                rubric = rubric_query
            
            # Generate document
            doc_buffer = self.generate_assignment_document(
                assignments=assignments,
                rubric=rubric,
                assignment_name=assignment_name
            )
            
            # Create filename
            safe_name = "".join(c for c in assignment_name if c.isalnum() or c in (' ', '-', '_')).strip()
            filename = f"{safe_name}_Assignment_Package"
            
            return self.create_download_response(doc_buffer, filename)
            
        except Exception as e:
            logger.error(f"Error generating assignments download: {str(e)}")
            raise

def generate_rubric_download(rubric_data: Dict) -> bytes:
    """Generate a Word document for a rubric"""
    
    # Create new document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add custom styles
    _add_custom_styles(doc)
    
    # Add rubric header
    rubric_name = rubric_data.get('rubric_name', 'Evaluation Rubric')
    doc.add_paragraph(rubric_name, style='Assignment Title')
    
    # Add document type
    doc_type = rubric_data.get('doc_type', 'Assignment')
    meta_para = doc.add_paragraph()
    meta_para.add_run("Document Type: ").bold = True
    meta_para.add_run(doc_type)
    
    # Add separator line
    doc.add_paragraph("_" * 80)
    
    # Add rubric categories
    rubrics = rubric_data.get('rubrics', [])
    for category_data in rubrics:
        # Category header
        category_para = doc.add_paragraph(style='Section Header')
        category_para.add_run(category_data.get('category', 'Category')).bold = True
        
        # Category questions
        questions = category_data.get('questions', [])
        for i, question in enumerate(questions, 1):
            question_para = doc.add_paragraph(style='Assignment Body')
            question_para.add_run(f"{i}. ").bold = True
            question_para.add_run(question)
        
        # Add spacing between categories
        doc.add_paragraph()
    
    # Save to BytesIO
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer.getvalue()

def _add_custom_styles(doc: Document):
    """Add custom styles to the document"""
    styles = doc.styles
    
    # Title style
    if 'Assignment Title' not in [s.name for s in styles]:
        title_style = styles.add_style('Assignment Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Calibri'
        title_font.size = Pt(18)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
    
    # Header style
    if 'Section Header' not in [s.name for s in styles]:
        header_style = styles.add_style('Section Header', WD_STYLE_TYPE.PARAGRAPH)
        header_font = header_style.font
        header_font.name = 'Calibri'
        header_font.size = Pt(12)
        header_font.bold = True
        header_style.paragraph_format.space_before = Pt(12)
        header_style.paragraph_format.space_after = Pt(6)
    
    # Body text style
    if 'Assignment Body' not in [s.name for s in styles]:
        body_style = styles.add_style('Assignment Body', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Calibri'
        body_font.size = Pt(11)
        body_style.paragraph_format.line_spacing = 1.15
        body_style.paragraph_format.space_after = Pt(6)
