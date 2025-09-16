"""
Text Extraction Service for extracting text from PDF and DOCX files
Handles both typed documents and coordinates with OCR service for handwritten content
"""

import os
import logging
import tempfile
from typing import Dict, Any, Optional
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document

from .ocr_service import OCRService

logger = logging.getLogger(__name__)

class TextExtractionService:
    """Service for extracting text from various document formats"""
    
    def __init__(self):
        """Initialize text extraction service"""
        self.ocr_service = OCRService()
        
    def extract_text_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF file, using OCR if necessary
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            logger.info(f"Extracting text from PDF: {file_path}")
            
            # First, try regular text extraction
            doc = fitz.open(file_path)
            extracted_text = ""
            total_pages = len(doc)
            
            for page_num in range(total_pages):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                extracted_text += f"\n\n--- Page {page_num + 1} ---\n{page_text}"
            
            doc.close()
            
            # Calculate text density to determine if OCR is needed
            avg_text_per_page = len(extracted_text.strip()) / total_pages if total_pages > 0 else 0
            
            if avg_text_per_page < 100:  # Likely handwritten/scanned
                logger.info(f"Low text density ({avg_text_per_page:.1f} chars/page), using OCR")
                
                # Check if OCR service is available
                ocr_status = self.ocr_service.test_connection()
                if ocr_status["status"] == "success":
                    ocr_result = self.ocr_service.transcribe_pdf(file_path)
                    return {
                        "text": ocr_result["full_text"],
                        "extraction_method": ocr_result.get("processing_method", "ocr"),
                        "confidence": ocr_result.get("ocr_confidence", 0.0),
                        "total_pages": total_pages,
                        "filename": Path(file_path).name,
                        "file_type": "pdf",
                        "ocr_details": ocr_result
                    }
                else:
                    logger.warning(f"OCR service unavailable: {ocr_status['message']}")
                    # Fall back to regular extraction even if text is sparse
                    return {
                        "text": extracted_text.strip(),
                        "extraction_method": "standard_fallback",
                        "confidence": 0.5,  # Lower confidence due to sparse text
                        "total_pages": total_pages,
                        "filename": Path(file_path).name,
                        "file_type": "pdf",
                        "warning": "OCR service unavailable, using standard extraction with low text density"
                    }
            else:
                logger.info(f"Good text density ({avg_text_per_page:.1f} chars/page), using standard extraction")
                return {
                    "text": extracted_text.strip(),
                    "extraction_method": "standard",
                    "confidence": 0.95,
                    "total_pages": total_pages,
                    "filename": Path(file_path).name,
                    "file_type": "pdf"
                }
                
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise Exception(f"PDF text extraction failed: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            logger.info(f"Extracting text from DOCX: {file_path}")
            
            doc = Document(file_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            extracted_text = "\n".join(full_text)
            
            return {
                "text": extracted_text,
                "extraction_method": "docx_standard",
                "confidence": 0.98,
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables),
                "filename": Path(file_path).name,
                "file_type": "docx"
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise Exception(f"DOCX text extraction failed: {str(e)}")
    
    def extract_text_from_bytes(self, file_bytes: bytes, filename: str, file_type: str) -> Dict[str, Any]:
        """
        Extract text from file bytes
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename
            file_type: File type (pdf or docx)
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            # Create temporary file
            suffix = f".{file_type.lower()}"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
                temp_file.write(file_bytes)
                temp_file_path = temp_file.name
            
            try:
                # Extract text based on file type
                if file_type.lower() == "pdf":
                    result = self.extract_text_from_pdf(temp_file_path)
                elif file_type.lower() == "docx":
                    result = self.extract_text_from_docx(temp_file_path)
                else:
                    raise Exception(f"Unsupported file type: {file_type}")
                
                # Update filename in result
                result["filename"] = filename
                return result
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            logger.error(f"Error extracting text from bytes ({filename}): {str(e)}")
            raise Exception(f"Text extraction failed: {str(e)}")
    
    def convert_docx_to_pdf_bytes(self, docx_bytes: bytes) -> bytes:
        """
        Convert DOCX bytes to PDF bytes for OCR processing if needed
        
        Args:
            docx_bytes: DOCX file content as bytes
            
        Returns:
            PDF file content as bytes
        """
        try:
            # This is a placeholder - in production, you might want to use
            # LibreOffice or another tool for high-fidelity conversion
            logger.warning("DOCX to PDF conversion not implemented - processing DOCX directly")
            raise NotImplementedError("DOCX to PDF conversion not available")
            
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {str(e)}")
            raise Exception(f"DOCX to PDF conversion failed: {str(e)}")
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from any supported file format
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == ".pdf":
                return self.extract_text_from_pdf(file_path)
            elif file_extension == ".docx":
                return self.extract_text_from_docx(file_path)
            else:
                raise Exception(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise Exception(f"Text extraction failed: {str(e)}")
    
    def get_extraction_summary(self, extraction_result: Dict[str, Any]) -> str:
        """
        Generate a human-readable summary of the extraction process
        
        Args:
            extraction_result: Result from text extraction
            
        Returns:
            Summary string
        """
        method = extraction_result.get("extraction_method", "unknown")
        confidence = extraction_result.get("confidence", 0.0)
        filename = extraction_result.get("filename", "unknown")
        text_length = len(extraction_result.get("text", ""))
        
        summary = f"Extracted {text_length} characters from {filename} using {method} "
        summary += f"(confidence: {confidence:.2f})"
        
        if "warning" in extraction_result:
            summary += f" - Warning: {extraction_result['warning']}"
        
        return summary
